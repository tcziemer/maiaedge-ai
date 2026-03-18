#!/usr/bin/env python3
"""
Build a MaiaEdge Proof of Concept Agreement from a template.

Takes a JSON input file with customer, equipment, and signer information,
loads the POC Word template, replaces variable fields using content-based
searching, populates tables, and saves the completed agreement.

Usage:
    python build_poc.py --input data.json [--template ../poc_agreement_template.docx]
"""

import argparse
import json
import sys
from pathlib import Path
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("ERROR: python-docx is not installed.", file=sys.stderr)
    print("Install it with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Input handling
# ---------------------------------------------------------------------------

def load_json_input(json_path: str) -> Dict[str, Any]:
    """Load and validate input data from JSON file."""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"ERROR: Input file not found: {json_path}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in {json_path}: {e}", file=sys.stderr)
        sys.exit(1)

    required = ['date', 'customer_name', 'customer_contact', 'poc_duration',
                 'equipment', 'delivery_locations', 'fees_description',
                 'maiaedge_signer_name', 'maiaedge_signer_title',
                 'customer_signer_name', 'customer_signer_title',
                 'white_logo_path', 'output_path']
    missing = [f for f in required if f not in data]
    if missing:
        print(f"ERROR: Missing required fields: {', '.join(missing)}", file=sys.stderr)
        sys.exit(1)

    if not data['equipment'] or not isinstance(data['equipment'], list):
        print("ERROR: 'equipment' must be a non-empty list", file=sys.stderr)
        sys.exit(1)

    if not data['delivery_locations'] or not isinstance(data['delivery_locations'], list):
        print("ERROR: 'delivery_locations' must be a non-empty list", file=sys.stderr)
        sys.exit(1)

    return data


# ---------------------------------------------------------------------------
# Robust paragraph-level replacement helpers
# ---------------------------------------------------------------------------

def replace_header_field(doc: Document, label: str, new_value: str) -> bool:
    """
    Replace a header field like 'Date:  March 6, 2026' or 'Customer:  Datum, Inc.'

    The label (e.g. 'Date:') may be split across multiple runs (bold label runs
    vs. normal value runs). This function works regardless of run boundaries by:
    1. Finding the paragraph whose text starts with the label
    2. Identifying which runs hold the "value" portion (after the label + whitespace)
    3. Clearing those value runs and setting the new value on the first one

    This preserves bold formatting on labels and normal formatting on values.
    """
    for para in doc.paragraphs:
        if not para.text.strip().startswith(label.rstrip(':')):
            continue

        runs = para.runs
        if not runs:
            continue

        # Walk through runs, accumulating text to find where the label ends
        # and the value begins
        accumulated = ""
        label_with_colon = label if label.endswith(':') else label + ':'
        value_start_idx = None

        for i, run in enumerate(runs):
            accumulated += run.text
            # Check if we've passed the label + colon + whitespace
            if value_start_idx is None and label_with_colon in accumulated:
                # The value starts either in this run (after the colon+space)
                # or in the next run
                after_label = accumulated.split(label_with_colon, 1)[1]
                if after_label.strip():
                    # Value is in this run (after the label)
                    value_start_idx = i
                    break
                else:
                    # Value starts in the next run
                    value_start_idx = i + 1
                    break

        if value_start_idx is None:
            # Label not found in runs — try a simpler approach
            value_start_idx = len(runs) - 1

        # Clear all value runs and set the first one
        for i in range(value_start_idx, len(runs)):
            if i == value_start_idx:
                # Check if this run also contains the colon
                if label_with_colon in runs[i].text:
                    runs[i].text = runs[i].text.split(label_with_colon)[0] + label_with_colon + "  " + new_value
                elif runs[i].text.strip() in (':', ':  ', ': '):
                    # This run is just the colon — value is in later runs
                    continue
                else:
                    runs[i].text = "  " + new_value if i > 0 else new_value
            else:
                runs[i].text = ""

        # Handle edge case: if value_start_idx == len(runs), the label+colon
        # was the entire paragraph text — need to update the last colon run
        if value_start_idx >= len(runs):
            # Find the run with the colon and append value
            for run in runs:
                if ':' in run.text:
                    run.text = run.text.rstrip() + "  " + new_value
                    break

        print(f"  ✓ Replaced {label} → {new_value}")
        return True

    print(f"  ⚠ WARNING: Could not find '{label}' field", file=sys.stderr)
    return False


def replace_text_in_cell(cell, old_text: str, new_text: str) -> bool:
    """Replace text within a table cell, preserving run formatting."""
    for para in cell.paragraphs:
        if old_text not in para.text:
            continue
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
        # If old_text spans runs, rebuild with full text replacement
        full = para.text.replace(old_text, new_text)
        if para.runs:
            para.runs[0].text = full
            for run in para.runs[1:]:
                run.text = ""
            return True
    return False


# ---------------------------------------------------------------------------
# Logo replacement
# ---------------------------------------------------------------------------

def replace_logo(doc: Document, white_logo_path: str) -> bool:
    """Replace the logo in Table 0, Row 0, Cell 0 with the white logo."""
    try:
        if len(doc.tables) == 0:
            print("  ⚠ WARNING: No tables found in document", file=sys.stderr)
            return False

        cell = doc.tables[0].rows[0].cells[0]

        # Remove all existing drawings (the old logo)
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for para in cell.paragraphs:
            for drawing in para._element.findall(f'.//{ns}drawing'):
                drawing.getparent().remove(drawing)
            for run in para.runs:
                run.text = ""

        if not Path(white_logo_path).exists():
            print(f"  ⚠ WARNING: Logo file not found: {white_logo_path}", file=sys.stderr)
            return False

        # Insert the new logo
        run = cell.paragraphs[0].add_run()
        run.add_picture(white_logo_path, width=Inches(2.5))
        print(f"  ✓ Replaced logo with white version")
        return True

    except Exception as e:
        print(f"  ⚠ WARNING: Failed to replace logo: {e}", file=sys.stderr)
        return False


# ---------------------------------------------------------------------------
# Table population
# ---------------------------------------------------------------------------

def populate_equipment_table(doc: Document, equipment: List[Dict]) -> bool:
    """Populate Equipment table (Table 1): SKU | Description | Qty | POC Fee"""
    try:
        if len(doc.tables) < 2:
            print("  ⚠ WARNING: Equipment table (Table 1) not found", file=sys.stderr)
            return False

        table = doc.tables[1]

        # Remove all data rows (keep header row 0)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        for item in equipment:
            row = table.add_row()
            row.cells[0].text = str(item.get('sku', ''))
            row.cells[1].text = str(item.get('description', ''))
            row.cells[2].text = str(item.get('qty', ''))
            row.cells[3].text = str(item.get('fee', ''))

        print(f"  ✓ Equipment table: {len(equipment)} rows added")
        return True

    except Exception as e:
        print(f"  ⚠ WARNING: Failed to populate equipment table: {e}", file=sys.stderr)
        return False


def populate_delivery_table(doc: Document, locations: List[Dict]) -> bool:
    """Populate Delivery Locations table (Table 2): # | Contact | Address | Notes"""
    try:
        if len(doc.tables) < 3:
            print("  ⚠ WARNING: Delivery table (Table 2) not found", file=sys.stderr)
            return False

        table = doc.tables[2]

        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        for item in locations:
            row = table.add_row()
            row.cells[0].text = str(item.get('number', ''))
            row.cells[1].text = str(item.get('contact', ''))
            row.cells[2].text = str(item.get('address', ''))
            row.cells[3].text = str(item.get('notes', ''))

        print(f"  ✓ Delivery table: {len(locations)} rows added")
        return True

    except Exception as e:
        print(f"  ⚠ WARNING: Failed to populate delivery table: {e}", file=sys.stderr)
        return False


# ---------------------------------------------------------------------------
# Section content and signature
# ---------------------------------------------------------------------------

def update_fees_section(doc: Document, fees_description: str) -> bool:
    """Update the fees description in Section 6 body text."""
    for para in doc.paragraphs:
        lower = para.text.lower()
        if ("at no charge" in lower or "complimentary" in lower) and "fee" in lower:
            for run in para.runs:
                if "no charge" in run.text.lower():
                    run.text = run.text.replace("no charge", fees_description)
                    print(f"  ✓ Updated fees: '{fees_description}'")
                    return True
                if "complimentary" in run.text.lower():
                    run.text = run.text.replace("complimentary", fees_description)
                    print(f"  ✓ Updated fees: '{fees_description}'")
                    return True

    print("  ⚠ Fees section not found or already correct — skipping", file=sys.stderr)
    return False


def update_customer_name_in_body(doc: Document, old_name: str, new_name: str) -> bool:
    """Replace the example customer name (e.g., 'Datum, Inc.') in section body text."""
    count = 0
    for para in doc.paragraphs:
        if old_name in para.text:
            for run in para.runs:
                if old_name in run.text:
                    run.text = run.text.replace(old_name, new_name)
                    count += 1
    if count > 0:
        print(f"  ✓ Replaced '{old_name}' → '{new_name}' in {count} locations")
    return count > 0


def update_signature_table(doc: Document, me_name: str, me_title: str,
                            cust_name: str, cust_title: str) -> bool:
    """Update the signature table (Table 3) with signer info."""
    try:
        if len(doc.tables) < 4:
            print("  ⚠ WARNING: Signature table (Table 3) not found", file=sys.stderr)
            return False

        sig = doc.tables[3]

        # MaiaEdge side (col 0)
        me_cell = sig.rows[0].cells[0]
        replace_text_in_cell(me_cell, "Tim Ziemer", me_name)
        # Handle title — look for the line after "Printed Name"
        for para in me_cell.paragraphs:
            for run in para.runs:
                if "CRO & Co-Founder" in run.text:
                    run.text = run.text.replace("CRO & Co-Founder", me_title)

        # Customer side (col 2)
        cust_cell = sig.rows[0].cells[2]
        # Replace the example customer name and title
        replace_text_in_cell(cust_cell, "Manish Singh", cust_name)
        # Find and replace the customer company name in the header
        replace_text_in_cell(cust_cell, "Datum, Inc.", cust_name.split(',')[0] if ',' in cust_name else cust_name)
        # Try to find/replace CTO or empty title
        for para in cust_cell.paragraphs:
            for run in para.runs:
                if "CTO" in run.text and run.text.strip() != "CTO":
                    run.text = run.text.replace("CTO", cust_title)
                elif run.text.strip() == "CTO":
                    run.text = run.text.replace("CTO", cust_title)

        print(f"  ✓ Signatures: MaiaEdge={me_name} ({me_title}), Customer={cust_name} ({cust_title})")
        return True

    except Exception as e:
        print(f"  ⚠ WARNING: Failed to update signature table: {e}", file=sys.stderr)
        return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Build a MaiaEdge POC Agreement from template")
    parser.add_argument('--input', required=True, help='JSON file with input data')
    parser.add_argument('--template', default='../poc_agreement_template.docx',
                        help='Path to POC template (default: ../poc_agreement_template.docx)')
    args = parser.parse_args()

    # Load input
    print("Loading input data...")
    data = load_json_input(args.input)
    print("✓ Input loaded\n")

    # Load template
    print(f"Loading template: {args.template}")
    try:
        doc = Document(args.template)
    except Exception as e:
        print(f"ERROR: Failed to load template: {e}", file=sys.stderr)
        sys.exit(1)
    print(f"✓ Template loaded ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)\n")

    # --- Header fields ---
    print("Header fields:")
    replace_header_field(doc, "Date:", data['date'])
    replace_header_field(doc, "Customer:", data['customer_name'])
    replace_header_field(doc, "Customer Contact:", data['customer_contact'])
    replace_header_field(doc, "POC Duration:", data['poc_duration'])

    # --- Logo ---
    print("\nLogo:")
    replace_logo(doc, data['white_logo_path'])

    # --- Tables ---
    print("\nTables:")
    populate_equipment_table(doc, data['equipment'])
    populate_delivery_table(doc, data['delivery_locations'])

    # --- Body text ---
    print("\nBody text:")
    update_customer_name_in_body(doc, "Datum, Inc.", data['customer_name'])
    update_fees_section(doc, data['fees_description'])

    # --- Signature ---
    print("\nSignature block:")
    update_signature_table(doc, data['maiaedge_signer_name'], data['maiaedge_signer_title'],
                            data['customer_signer_name'], data['customer_signer_title'])

    # --- Save ---
    output_path = data['output_path']
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"\n✓ POC Agreement saved to: {output_path}")


if __name__ == '__main__':
    main()
