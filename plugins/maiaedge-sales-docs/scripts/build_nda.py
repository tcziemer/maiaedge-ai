#!/usr/bin/env python3
"""
Build a MaiaEdge Mutual NDA from a template.

This script takes a JSON input file with participant and signer information,
loads a Word document template, replaces variable fields using content-based searching,
and saves the completed NDA.

Usage:
    python build_nda.py --input data.json [--template ../nda_template.docx] [--output output.docx]
"""

import argparse
import json
import sys
from pathlib import Path
from typing import Dict, Any

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("ERROR: python-docx is not installed.", file=sys.stderr)
    print("Install it with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def load_json_input(json_path: str) -> Dict[str, Any]:
    """Load input data from JSON file."""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except FileNotFoundError:
        print(f"ERROR: Input file not found: {json_path}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in {json_path}: {e}", file=sys.stderr)
        sys.exit(1)


def validate_input_data(data: Dict[str, Any]) -> None:
    """Validate that all required fields are present."""
    required_fields = [
        'effective_date',
        'participant_name',
        'participant_address',
        'signer_name',
        'signer_title',
        'output_path'
    ]

    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        print(f"ERROR: Missing required fields in JSON: {', '.join(missing_fields)}", file=sys.stderr)
        sys.exit(1)


def load_template(template_path: str) -> Document:
    """Load the NDA template document."""
    try:
        doc = Document(template_path)
        return doc
    except FileNotFoundError:
        print(f"ERROR: Template file not found: {template_path}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: Failed to load template: {e}", file=sys.stderr)
        sys.exit(1)


def replace_date_in_preamble(doc: Document, effective_date: str) -> bool:
    """
    Replace the date blank in the preamble paragraph.

    Searches for the paragraph containing "________, 202_" and replaces it
    with the effective date.
    """
    for para in doc.paragraphs:
        if "________, 202_" in para.text:
            print(f"  Found date placeholder in paragraph: {para.text[:50]}...")

            # Find and replace the underscore blanks
            for run in para.runs:
                if "________, 202_" in run.text:
                    run.text = run.text.replace("________, 202_", effective_date)
                    print(f"  ✓ Replaced date with: {effective_date}")
                    return True

            # If not found in a single run, reconstruct from multiple runs
            para_text = para.text
            if "________" in para_text and "202_" in para_text:
                # Rebuild the paragraph text by replacing across runs
                full_text = ''.join([run.text for run in para.runs])
                new_text = full_text.replace("________, 202_", effective_date)

                # Clear existing runs and create new one with replacement
                for run in para.runs:
                    run.text = ""
                para.add_run(new_text)
                print(f"  ✓ Replaced date (multi-run) with: {effective_date}")
                return True

    print("  ⚠ WARNING: Could not find date placeholder", file=sys.stderr)
    return False


def replace_participant_name(doc: Document, participant_name: str) -> bool:
    """
    Replace the participant name blank.

    Searches for the paragraph containing "Participant:" and replaces
    the blanks on the right side with participant_name.
    """
    for para in doc.paragraphs:
        if "Participant:" in para.text:
            print(f"  Found participant placeholder in paragraph: {para.text[:60]}...")

            # The pattern is "Participant:\t_______________________________"
            # We need to replace the blanks after the tab
            para_text = para.text

            # Replace the blanks that come after "Participant:"
            if "Participant:" in para_text:
                # Find the position of "Participant:" and replace the blanks after it
                parts = para_text.split("Participant:")
                if len(parts) >= 2:
                    after_label = parts[1]
                    # Replace leading blanks with the name
                    new_after = after_label.lstrip('_\t ').lstrip()

                    # Reconstruct by replacing underscore blanks
                    for run in para.runs:
                        if "_______________________________" in run.text and "Participant:" not in run.text:
                            run.text = run.text.replace("_______________________________", participant_name)
                            print(f"  ✓ Replaced participant name with: {participant_name}")
                            return True

            # Fallback: try to replace underscores in the run after "Participant:"
            found_label = False
            for run in para.runs:
                if found_label and "_" in run.text:
                    run.text = run.text.replace("_" * len(run.text), participant_name)
                    print(f"  ✓ Replaced participant name with: {participant_name}")
                    return True
                if "Participant:" in run.text:
                    found_label = True

    print("  ⚠ WARNING: Could not find participant name placeholder", file=sys.stderr)
    return False


def replace_participant_address(doc: Document, participant_address: str) -> bool:
    """
    Replace the participant address blanks.

    Searches for the paragraph containing "Address:" and replaces the blanks
    with participant_address. Also handles the continuation line.
    """
    address_replaced = False

    for i, para in enumerate(doc.paragraphs):
        if "Address:" in para.text:
            print(f"  Found address placeholder in paragraph {i}: {para.text[:60]}...")

            # Replace blanks after "Address:" label in this paragraph
            for run in para.runs:
                if "_______________________________" in run.text and "Address:" not in run.text:
                    run.text = run.text.replace("_______________________________", participant_address)
                    print(f"  ✓ Replaced first address line with: {participant_address}")
                    address_replaced = True
                    break

            # Check next paragraph for continuation address blanks
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                # Look for continuation blanks (usually starts with blank lines or tabs)
                if "_______________________________" in next_para.text and "MaiaEdge" not in next_para.text:
                    for run in next_para.runs:
                        if "_______________________________" in run.text:
                            run.text = run.text.replace("_______________________________", "")
                            print(f"  ✓ Cleared continuation address line")
                            address_replaced = True
                            break

            if address_replaced:
                return True

    if not address_replaced:
        print("  ⚠ WARNING: Could not find address placeholder", file=sys.stderr)

    return address_replaced


def replace_signer_name(doc: Document, signer_name: str) -> bool:
    """
    Replace the signer name blank.

    Searches for the paragraph containing "Printed Name:" and replaces
    the SECOND set of blanks (right side) with signer_name.
    """
    for para in doc.paragraphs:
        if "Printed Name:" in para.text:
            print(f"  Found signer name placeholder in paragraph: {para.text[:70]}...")

            # The pattern has "Printed Name:" twice (left and right sides)
            # We need to replace the SECOND set of blanks (right side)
            para_text = para.text
            blank_count = 0

            for run in para.runs:
                if "_______________________________" in run.text:
                    blank_count += 1
                    # Replace the second occurrence (right side)
                    if blank_count == 2:
                        run.text = run.text.replace("_______________________________", signer_name)
                        print(f"  ✓ Replaced signer name with: {signer_name}")
                        return True

    print("  ⚠ WARNING: Could not find signer name placeholder", file=sys.stderr)
    return False


def replace_signer_title(doc: Document, signer_title: str) -> bool:
    """
    Replace the signer title blank.

    Searches for the paragraph containing "Title:" and replaces
    the SECOND set of blanks (right side) with signer_title.
    """
    for para in doc.paragraphs:
        if "Title:" in para.text:
            print(f"  Found signer title placeholder in paragraph: {para.text[:70]}...")

            # The pattern has "Title:" twice (left and right sides)
            # We need to replace the SECOND set of blanks (right side)
            para_text = para.text
            blank_count = 0

            for run in para.runs:
                if "_______________________________" in run.text:
                    blank_count += 1
                    # Replace the second occurrence (right side)
                    if blank_count == 2:
                        run.text = run.text.replace("_______________________________", signer_title)
                        print(f"  ✓ Replaced signer title with: {signer_title}")
                        return True

    print("  ⚠ WARNING: Could not find signer title placeholder", file=sys.stderr)
    return False


def save_document(doc: Document, output_path: str) -> bool:
    """Save the modified document to output path."""
    try:
        doc.save(output_path)
        print(f"✓ Document saved to: {output_path}")
        return True
    except Exception as e:
        print(f"ERROR: Failed to save document: {e}", file=sys.stderr)
        return False


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Build a MaiaEdge Mutual NDA from a template",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python build_nda.py --input data.json
  python build_nda.py --input data.json --template /path/to/nda_template.docx
  python build_nda.py --input data.json --template ../nda_template.docx
        """
    )

    parser.add_argument(
        '--input',
        required=True,
        help='JSON file with input data (effective_date, participant_name, participant_address, signer_name, signer_title, output_path)'
    )

    parser.add_argument(
        '--template',
        default='../nda_template.docx',
        help='Path to NDA template DOCX file (default: ../nda_template.docx)'
    )

    args = parser.parse_args()

    # Load and validate input
    print("Loading input data...")
    input_data = load_json_input(args.input)
    validate_input_data(input_data)
    print(f"✓ Input loaded successfully")

    # Load template
    print(f"\nLoading template from: {args.template}")
    doc = load_template(args.template)
    print(f"✓ Template loaded ({len(doc.paragraphs)} paragraphs)")

    # Perform replacements
    print("\nReplacing variables:")

    print("  • Effective Date:")
    if not replace_date_in_preamble(doc, input_data['effective_date']):
        sys.exit(1)

    print("  • Participant Name:")
    if not replace_participant_name(doc, input_data['participant_name']):
        sys.exit(1)

    print("  • Participant Address:")
    if not replace_participant_address(doc, input_data['participant_address']):
        sys.exit(1)

    print("  • Signer Name:")
    if not replace_signer_name(doc, input_data['signer_name']):
        sys.exit(1)

    print("  • Signer Title:")
    if not replace_signer_title(doc, input_data['signer_title']):
        sys.exit(1)

    # Save output
    print(f"\nSaving document:")
    output_path = input_data['output_path']

    # Create output directory if needed
    output_dir = Path(output_path).parent
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        print(f"ERROR: Failed to create output directory: {e}", file=sys.stderr)
        sys.exit(1)

    if not save_document(doc, output_path):
        sys.exit(1)

    print("\n✓ NDA generation completed successfully")
    sys.exit(0)


if __name__ == '__main__':
    main()
